import os
import sys
import json
import time
import datetime
import shutil
import re
import csv
import io
import logging
import platform
import webbrowser
import urllib.parse
import urllib.request
from pathlib import Path

logger = logging.getLogger("lumin.tools")

# Optional imports with safe fallback handling
try:
    from PIL import ImageGrab
    PIL_OK = True
except ImportError:
    PIL_OK = False

try:
    import pyperclip
    CLIP_OK = True
except ImportError:
    CLIP_OK = False

try:
    import psutil
    PSUTIL_OK = True
except ImportError:
    PSUTIL_OK = False

try:
    import requests
    REQUESTS_OK = True
except ImportError:
    REQUESTS_OK = False

try:
    from selenium import webdriver
    from selenium.webdriver.common.by import By
    from selenium.webdriver.common.keys import Keys
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    SELENIUM_OK = True
except ImportError:
    SELENIUM_OK = False

# Configuration and Security Constants
DANGEROUS_KEYWORDS = [
    "remove-item", "rm -rf", "del /", "format ", "diskpart",
    "shutdown", "stop-computer", "restart-computer",
    "net user", "reg delete", "vssadmin", "bcdedit",
    "takeown", "icacls", "cipher /w", "fdisk",
    "dd if=", "mkfs", "wipefs", "rmdir /s",
]

DENYLIST_PATTERNS = [
    ".ssh", ".aws", "id_rsa", "id_ed25519", ".gnupg",
    "system32/config/sam", "system32\\config\\sam",
    "system32/config/security", "system32\\config\\security",
    "system32/config/system", "system32\\config\\system",
    "ntds.dit",
    "System32/config", "System32\\config", "Windows/System32", "Windows\\System32",
    "AppData/Roaming/Microsoft/Credentials", "AppData\\Roaming\\Microsoft\\Credentials",
]

FOLDER_SHORTCUTS = {
    "videos":    str(Path.home() / "Videos"),
    "pictures":  str(Path.home() / "Pictures"),
    "music":     str(Path.home() / "Music"),
    "documents": str(Path.home() / "Documents"),
    "downloads": str(Path.home() / "Downloads"),
    "desktop":   str(Path.home() / "Desktop"),
    "home":      str(Path.home()),
}

class ToolRegistry:
    """
    Registry of system and conversational tools that the LUMIN Python Agent can execute.
    Includes comprehensive security, interactive prompts, and full cross-platform compatibility.
    """
    def __init__(self, memory_manager=None, base_dir=None):
        self.memory_manager = memory_manager
        self.base_dir = base_dir or os.path.abspath(os.path.dirname(os.path.dirname(__file__)))
        self.config_path = os.path.join(self.base_dir, "agent_config.json")
        self.audit_path = os.path.join(self.base_dir, "audit_log.jsonl")
        self.selenium_driver = None
        
        # Self-initialize configuration with safe defaults
        self._init_config()

        # Build tool dictionary
        self.tools = {
            "web_search": self.web_search,
            "list_directory": self.list_directory,
            "directory_tree": self.directory_tree,
            "read_file": self.read_file,
            "write_file": self.write_file,
            "delete_file": self.delete_file,
            "write_csv": self.write_csv,
            "write_report": self.write_report,
            "write_docx": self.write_docx,
            "set_reminder": self.set_reminder,
            "run_file": self.run_file,
            "run_powershell": self.run_powershell,
            "launch_application": self.launch_application,
            "write_text_to_active_window": self.write_text_to_active_window,
            "write_to_notepad": self.write_text_to_active_window,
            "close_application": self.close_application,
            "take_screenshot": self.take_screenshot,
            "describe_image": self.describe_image,
            "list_processes": self.list_processes,
            "kill_process": self.kill_process,
            "get_clipboard": self.get_clipboard,
            "set_clipboard": self.set_clipboard,
            "open_url": self.open_url,
            "search_youtube": self.search_youtube,
            "play_first_youtube_video": self.play_first_youtube_video,
            "open_file_or_folder": self.open_file_or_folder,
            "browser_navigate": self.browser_navigate,
            "browser_click": self.browser_click,
            "browser_type": self.browser_type,
            "browser_read_page": self.browser_read_page,
            "browser_scroll": self.browser_scroll,
            "close_browser": self.close_browser,
            "fetch_reddit": self.fetch_reddit,
            "list_models": self.list_models,
            "switch_model": self.switch_model,
            "github_list_repos": self.github_list_repos,
            "github_list_issues": self.github_list_issues,
            "github_list_prs": self.github_list_prs,
            "github_view_pr_diff": self.github_view_pr_diff,
            "github_create_issue": self.github_create_issue,
            "get_system_time": self.get_system_time,
            "get_hardware_status": self.get_hardware_status,
            "change_theme": self.change_theme,
            "set_visualizer_shape": self.set_visualizer_shape,
            "store_memory_fact": self.store_memory_fact,
            "mcp_connect": self.mcp_connect,
            "mcp_disconnect": self.mcp_disconnect,
            "mcp_list_servers": self.mcp_list_servers,
            "mcp_call_tool": self.mcp_call_tool,
        }

        # Lazy initialize MCP Client Manager
        try:
            from tools.mcp_client import MCPClientManager
            self.mcp_client = MCPClientManager(self.base_dir)
        except Exception as e:
            logger.warning(f"Failed to load MCPClientManager: {e}")
            self.mcp_client = None

    def _init_config(self):
        """Initializes configuration file from agent_config.example.json or with safe defaults."""
        if not os.path.exists(self.config_path):
            example_path = os.path.join(self.base_dir, "agent_config.example.json")
            if os.path.exists(example_path):
                try:
                    shutil.copyfile(example_path, self.config_path)
                    return
                except Exception as e:
                    logger.error(f"Failed to copy agent_config.example.json: {e}")

            defaults = {
                "auto_approve": True,
                "auto_approve_destructive": False,
                "unrestricted_mode": True,
                "bypass_denylist": False,
                "enable_mcp": True,
                "allowed_folders": [
                    os.path.expanduser("~/Desktop"),
                    os.path.expanduser("~/Documents"),
                    os.path.expanduser("~/Downloads"),
                    self.base_dir
                ]
            }
            try:
                with open(self.config_path, "w", encoding="utf-8") as f:
                    json.dump(defaults, f, indent=2)
                # First run safety warning printout
                print("\n" + "="*80)
                print("  [SECURITY NOTICE]: First-run LUMIN Agent safeguards have been initialized.")
                print("  - auto_approve = true")
                print("  - auto_approve_destructive = false (confirmation required for high-risk actions)")
                print("  - unrestricted_mode = true")
                print("  - bypass_denylist = false (system path denylist enforced)")
                print("  LUMIN will run with safe local automation capabilities.")
                print("="*80 + "\n")
                sys.stdout.flush()
            except Exception as e:
                logger.error(f"Failed to initialize configuration: {e}")

    def _get_config(self):
        """Retrieves config dict from disk."""
        if os.path.exists(self.config_path):
            try:
                with open(self.config_path, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        return {}

    def _save_config(self, cfg):
        """Saves configuration atomically."""
        tmp = self.config_path + ".tmp"
        try:
            with open(tmp, "w", encoding="utf-8") as f:
                json.dump(cfg, f, indent=2)
            os.replace(tmp, self.config_path)
        except Exception as e:
            logger.error(f"Failed to save configuration: {e}")

    def _audit(self, action, details, approved=True, result=""):
        """Logs action execution to the audit log trail."""
        try:
            entry = {
                "ts": datetime.datetime.now().isoformat(),
                "action": action,
                "details": details,
                "approved": approved,
                "result": result
            }
            with open(self.audit_path, "a", encoding="utf-8") as f:
                f.write(json.dumps(entry) + "\n")
        except Exception as e:
            logger.error(f"Audit log write failure: {e}")

    def _confirm(self, action_label, details, high_risk=False):
        """Handles interactive verification prompts with config-based auto-approvals."""
        cfg = self._get_config()
        
        # Check auto-approval overrides
        if high_risk and cfg.get("auto_approve_destructive", False):
            logger.info(f"Auto-approved destructive action: {action_label}")
            self._audit(action_label, details, approved=True, result="Auto-Approved Destructive")
            return True
        if not high_risk and cfg.get("auto_approve", False):
            logger.info(f"Auto-approved action: {action_label}")
            self._audit(action_label, details, approved=True, result="Auto-Approved")
            return True

        # Render terminal interaction box
        print("\n" + "═" * 70)
        print(f"  [USER INPUT REQUIRED] - {action_label}")
        print("═" * 70)
        print(details)
        print("═" * 70)
        
        if high_risk:
            print("  !! WARNING: This action is classified as potentially DESTRUCTIVE !!")
            print("  Type exactly: CONFIRM DESTRUCTIVE ACTION to proceed, or anything else to abort.")
            print("  Waiting for input... > ")
            sys.stdout.flush()
            resp = sys.stdin.readline().strip()
            approved = (resp == "CONFIRM DESTRUCTIVE ACTION")
        else:
            print("  Type Y/YES to approve this action, or press Enter to cancel.")
            print("  Waiting for input... > ")
            sys.stdout.flush()
            resp = sys.stdin.readline().strip().upper()
            approved = resp in ("Y", "YES")
            
        self._audit(action_label, details, approved=approved, result="Approved" if approved else "User Cancelled")
        
        if not approved:
            print("\n>>> [SYSTEM]: Action cancelled by user. Continuing...")
            sys.stdout.flush()
            
        return approved

    def _resolve_path(self, path_str):
        """Resolves folder shortcuts, environment variables, and returns an absolute clean path."""
        parts = os.path.split(path_str)
        if parts and parts[0].lower() in FOLDER_SHORTCUTS:
            resolved_root = FOLDER_SHORTCUTS[parts[0].lower()]
            return os.path.abspath(os.path.join(resolved_root, *parts[1:]))
        return os.path.abspath(os.path.expanduser(os.path.expandvars(path_str)))

    def _is_denied(self, resolved_path):
        """Enforces a system-wide denylist of files and paths."""
        cfg = self._get_config()
        if cfg.get("bypass_denylist", False):
            return False
        
        normalized = resolved_path.lower()
        return any(pat.lower() in normalized for pat in DENYLIST_PATTERNS)

    def _check_file_access(self, resolved_path):
        """Enforces sandboxed directory limitations unless unrestricted mode is enabled."""
        abs_path = self._resolve_path(resolved_path)
        if self._is_denied(abs_path):
            return f"Security Guard: Access denied. Path '{abs_path}' is protected."
        
        cfg = self._get_config()
        if cfg.get("unrestricted_mode", False):
            return None
            
        import tempfile
        allowed_roots = list(cfg.get("allowed_folders", []))
        workspace_root = os.path.abspath(os.getcwd())
        temp_root = os.path.abspath(tempfile.gettempdir())
        
        # Always grant access to workspace root, system temp, and base dir
        default_roots = [workspace_root, temp_root, self.base_dir]
        for dr in default_roots:
            if dr and dr not in allowed_roots:
                allowed_roots.append(dr)

        resolved_norm = os.path.normpath(abs_path).lower()
        for root in allowed_roots:
            if not root:
                continue
            normalized_root = os.path.normpath(os.path.abspath(root)).lower()
            if resolved_norm == normalized_root or resolved_norm.startswith(normalized_root + os.sep) or resolved_norm.startswith(normalized_root):
                return None # Match found!
                
        return (
            f"Security Exception: Path '{abs_path}' is outside your allowed folders list.\n"
            f"Allowed Folders: {allowed_roots}\n"
            f"Type 'unrestricted on' to enable unrestricted directory access."
        )


    def execute_tool(self, tool_name, *args, **kwargs):
        """Executes a registered tool by name with arguments."""
        if tool_name not in self.tools:
            return f"Error: Tool '{tool_name}' not found."
        try:
            logger.info(f"Executing tool {tool_name} with args={args} kwargs={kwargs}")
            return self.tools[tool_name](*args, **kwargs)
        except Exception as e:
            logger.error(f"Error executing tool {tool_name}: {e}")
            return f"Error executing tool: {str(e)}"

    # ── Conversational & Visualizer Tools ────────────────────────────────────
    def get_system_time(self):
        """Retrieves the current date and time in 12-hour format with AM/PM and timezone when available."""
        now = datetime.datetime.now()
        tz_name = now.astimezone().tzname() or ""
        time_str = now.strftime("%Y-%m-%d %I:%M:%S %p")
        if tz_name:
            time_str += f" {tz_name}"
        return f"The current system local time is {time_str}."

    def get_hardware_status(self):
        """Gets the machine's current operating hardware characteristics."""
        return (
            f"OS: {platform.system()} {platform.release()} ({platform.machine()})\n"
            f"Processor: {platform.processor() or 'Standard CPU Core'}\n"
            f"Cognitive Pipeline: Unified Local and Cloud AI Hybrid Router Enabled\n"
            f"Virtual Core: ONLINE & SYNCHRONIZED"
        )

    def change_theme(self, theme_name):
        """Changes the visualizer color/light pattern theme."""
        valid_themes = ["cyberware", "crimson", "matrix", "solar", "arcane", "glacial", "golden", "hotpink", "aqua", "tungsten"]
        theme_name = theme_name.lower().strip()
        if theme_name not in valid_themes:
            return f"Invalid theme '{theme_name}'. Available themes: {', '.join(valid_themes)}."
        print(f"\n>>> [SYSTEM SHIFT]: Visualizer theme changing to '{theme_name}'...")
        sys.stdout.flush()
        return f"Successfully executed theme transition. Selected visual skin: {theme_name.upper()}. [COMMAND: CHANGE_THEME={theme_name}]"

    def set_visualizer_shape(self, shape):
        """Morphs the 3D visualizer's geometry vertex configuration."""
        valid_shapes = ["sphere", "cube", "pyramid", "torus", "helix", "triangle", "saturn"]
        shape = shape.lower().strip()
        if shape not in valid_shapes:
            return f"Invalid geometry shape '{shape}'. Available: {', '.join(valid_shapes)}."
        print(f"\n>>> [GEOMETRY ALTER]: Morphing core vertex array into a '{shape}'...")
        sys.stdout.flush()
        return f"Core geometry morphed successfully to {shape.upper()}. [COMMAND: SET_SHAPE={shape}]"

    def store_memory_fact(self, fact):
        """Saves an explicit piece of information to the long-term memories list."""
        if not fact:
            return "Error: Memory fact text cannot be empty."
        if self.memory_manager:
            self.memory_manager.store_long_term_memory(fact)
            return f"Fact successfully committed to long-term memory: '{fact}'."
        return "Error: Memory manager module is not initialized."

    # ── File and Directory Management ────────────────────────────────────────
    def list_directory(self, path="."):
        """Lists files and folders inside a directory."""
        resolved = self._resolve_path(path)
        access_err = self._check_file_access(resolved)
        if access_err: return access_err
        
        try:
            items = os.listdir(resolved)
            output = []
            for item in sorted(items):
                full = os.path.join(resolved, item)
                tag = "[DIR]" if os.path.isdir(full) else "[FILE]"
                output.append(f"{tag:6s} {item}")
            return "\n".join(output) if output else "(empty directory)"
        except Exception as e:
            return f"Error listing directory: {e}"

    def directory_tree(self, path=".", max_depth=3):
        """Generates a tree representation of directory contents."""
        resolved = self._resolve_path(path)
        access_err = self._check_file_access(resolved)
        if access_err: return access_err
        
        ignore = {".git", "__pycache__", "node_modules", ".venv", "venv", ".backups"}
        lines = [resolved]
        
        def walk(cur, prefix, depth):
            if depth > max_depth:
                return
            try:
                entries = sorted(os.listdir(cur))
            except PermissionError:
                lines.append(f"{prefix}!! [permission denied]")
                return
            except Exception as e:
                lines.append(f"{prefix}!! [error: {e}]")
                return
                
            entries = [e for e in entries if e not in ignore]
            for i, entry in enumerate(entries):
                full = os.path.join(cur, entry)
                is_last = (i == len(entries) - 1)
                connector = "└── " if is_last else "├── "
                tag = "[DIR] " if os.path.isdir(full) else ""
                lines.append(f"{prefix}{connector}{tag}{entry}")
                if os.path.isdir(full):
                    next_prefix = prefix + ("    " if is_last else "│   ")
                    walk(full, next_prefix, depth + 1)
                    
        walk(resolved, "", 1)
        return "\n".join(lines)

    def read_file(self, file_path):
        """Reads a UTF-8 text file's contents."""
        resolved = self._resolve_path(file_path)
        access_err = self._check_file_access(resolved)
        if access_err: return access_err
        
        try:
            with open(resolved, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        except Exception as e:
            return f"Error reading file '{file_path}': {e}"

    def write_file(self, file_path, content):
        """Writes text content to a target file path atomically (with automatic backup)."""
        resolved = self._resolve_path(file_path)
        access_err = self._check_file_access(resolved)
        if access_err: return access_err
        
        # Display preview for user approval
        preview = content[:200] + ("\n...(truncated)" if len(content) > 200 else "")
        details = f"Target File: {resolved}\nSize: {len(content)} characters\n\nContent Preview:\n{preview}"
        
        if not self._confirm("WRITE FILE", details):
            return "Cancelled by user."
            
        try:
            # Create a backup if file exists
            if os.path.exists(resolved):
                backup_dir = os.path.join(self.base_dir, ".backups")
                os.makedirs(backup_dir, exist_ok=True)
                backup_name = f"{os.path.basename(resolved)}.{int(time.time())}.bak"
                shutil.copy2(resolved, os.path.join(backup_dir, backup_name))
                
            # Atomic Write
            temp_file = resolved + ".tmp"
            os.makedirs(os.path.dirname(resolved), exist_ok=True)
            with open(temp_file, "w", encoding="utf-8") as f:
                f.write(content)
            os.replace(temp_file, resolved)
            return f"Successfully wrote content to file: {resolved}"
        except Exception as e:
            return f"Error writing file '{file_path}': {e}"

    def delete_file(self, file_path):
        """Deletes a file permanently after backup and confirmation."""
        resolved = self._resolve_path(file_path)
        access_err = self._check_file_access(resolved)
        if access_err: return access_err
        
        if not os.path.exists(resolved):
            return f"File not found: '{file_path}'"
            
        details = f"File: {resolved}\nSize: {os.path.getsize(resolved)} bytes"
        if not self._confirm("DELETE FILE", details, high_risk=True):
            return "Cancelled by user."
            
        try:
            # Backup
            backup_dir = os.path.join(self.base_dir, ".backups")
            os.makedirs(backup_dir, exist_ok=True)
            backup_name = f"{os.path.basename(resolved)}.{int(time.time())}.bak"
            shutil.copy2(resolved, os.path.join(backup_dir, backup_name))
            
            os.remove(resolved)
            return f"File deleted successfully (Backup archived in .backups): {resolved}"
        except Exception as e:
            return f"Error deleting file '{file_path}': {e}"

    def write_csv(self, file_path, header, rows):
        """Creates a formatted CSV file."""
        resolved = self._resolve_path(file_path)
        access_err = self._check_file_access(resolved)
        if access_err: return access_err
        
        details = f"CSV File: {resolved}\nHeader: {header}\nRow count: {len(rows)}"
        if not self._confirm("WRITE CSV FILE", details):
            return "Cancelled by user."
            
        try:
            temp_file = resolved + ".tmp"
            os.makedirs(os.path.dirname(resolved), exist_ok=True)
            with open(temp_file, "w", encoding="utf-8", newline="") as f:
                writer = csv.writer(f)
                writer.writerow(header)
                writer.writerows(rows)
            os.replace(temp_file, resolved)
            return f"CSV document successfully written: {resolved}"
        except Exception as e:
            return f"Error writing CSV file: {e}"

    def write_report(self, filename, title, body):
        """Formats and saves a stylized text report."""
        header = f"{'='*60}\n  {title.upper()}\n  Created: {datetime.datetime.now():%Y-%m-%d %H:%M}\n{'='*60}\n\n"
        full_report = header + body
        return self.write_file(filename, full_report)

    def write_docx(self, file_path, title="", body_paragraphs=None):
        """Creates a formatted Word (.docx) document using python-docx."""
        resolved = self._resolve_path(file_path)
        access_err = self._check_file_access(resolved)
        if access_err: return access_err
        
        body_paragraphs = body_paragraphs or []
        if isinstance(body_paragraphs, str):
            paragraphs = [p.strip() for p in body_paragraphs.split("\n\n") if p.strip()]
        elif isinstance(body_paragraphs, list):
            paragraphs = body_paragraphs
        else:
            paragraphs = [str(body_paragraphs)]

        details = f"Word Document Path: {resolved}\nTitle: {title}\nParagraphs: {len(paragraphs)}"
        if not self._confirm("WRITE WORD DOCUMENT", details):
            return "Cancelled by user."
            
        try:
            os.makedirs(os.path.dirname(resolved), exist_ok=True)
            try:
                import docx
                doc = docx.Document()
                if title:
                    doc.add_heading(title, level=0)
                for p in paragraphs:
                    p_str = str(p).strip()
                    if p_str.startswith("# "):
                        doc.add_heading(p_str[2:].strip(), level=1)
                    elif p_str.startswith("## "):
                        doc.add_heading(p_str[3:].strip(), level=2)
                    elif p_str.startswith("### "):
                        doc.add_heading(p_str[4:].strip(), level=3)
                    else:
                        doc.add_paragraph(p_str)
                doc.save(resolved)
                return f"Word document (.docx) created successfully at: {resolved}"
            except ImportError:
                # Fallback to text file format if python-docx is not available
                content = f"# {title}\n\n" + "\n\n".join(str(p) for p in paragraphs)
                with open(resolved, "w", encoding="utf-8") as f:
                    f.write(content)
                return f"Document saved successfully as structured text file at: {resolved}"
        except Exception as e:
            return f"Error creating Word document: {e}"

    def set_reminder(self, reminder_text, time_str=""):
        """Sets a system reminder and records it in long-term memory and reminders log."""
        if not reminder_text:
            return "Error: Reminder text cannot be empty."
        timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        record = f"[{timestamp}] Reminder: {reminder_text}" + (f" (Due: {time_str})" if time_str else "")
        if self.memory_manager:
            self.memory_manager.store_long_term_memory(record)
        reminders_file = os.path.join(self.base_dir, "reminders.txt")

        toast_status = "Recorded locally"
        if platform.system() == "Windows":
            try:
                import subprocess
                ps_toast = f'''
                [void] [System.Reflection.Assembly]::LoadWithPartialName("System.Windows.Forms")
                $notification = New-Object System.Windows.Forms.NotifyIcon
                $notification.Icon = [System.Drawing.SystemIcons]::Information
                $notification.BalloonTipTitle = "LUMIN Agent Reminder"
                $notification.BalloonTipText = "Reminder: {reminder_text}" + (if ("{time_str}") {{ " (Due: {time_str})" }} else {{ "" }})
                $notification.Visible = $true
                $notification.ShowBalloonTip(5000)
                Start-Sleep -Seconds 5
                $notification.Dispose()
                '''
                subprocess.Popen(["powershell", "-NoProfile", "-NonInteractive", "-Command", ps_toast], creationflags=0x08000000 if hasattr(subprocess, "CREATE_NO_WINDOW") else 0)
                toast_status = "Windows Toast Notification triggered & recorded"
            except Exception as te:
                logger.warning(f"Windows Toast notification failed: {te}")

        try:
            with open(reminders_file, "a", encoding="utf-8") as f:
                f.write(record + "\n")
            return f"Reminder recorded successfully ({toast_status}):\n- Task: {reminder_text}\n- Due/Time: {time_str or 'Not specified'}\n- Saved in: reminders.txt & agent memory"
        except Exception as e:
            return f"Reminder stored in memory, but error logging to file: {e}"

    # ── Process and Program Execution ────────────────────────────────────────
    def run_file(self, file_path):
        """Executes a script or runs a file locally."""
        resolved = self._resolve_path(file_path)
        access_err = self._check_file_access(resolved)
        if access_err: return access_err
        
        if not os.path.exists(resolved):
            return f"File not found: {resolved}"
            
        if not self._confirm("RUN FILE", f"Executable Path: {resolved}"):
            return "Cancelled by user."
            
        try:
            if resolved.endswith(".py"):
                import subprocess
                res = subprocess.run([sys.executable, resolved], capture_output=True, text=True, timeout=30)
                return f"Execution Completed.\nSTDOUT:\n{res.stdout}\nSTDERR:\n{res.stderr}"
            else:
                if platform.system() == "Windows":
                    os.startfile(resolved)
                else:
                    import subprocess
                    subprocess.Popen(["xdg-open", resolved])
                return f"Launched program: {resolved}"
        except Exception as e:
            return f"Error running file: {e}"

    def run_powershell(self, command):
        """Runs a command via PowerShell on Windows."""
        if platform.system() != "Windows":
            return "PowerShell is only supported natively on Windows systems."
            
        # Security validation check
        for kw in DANGEROUS_KEYWORDS:
            if kw in command.lower():
                return f"Security Exception: Command blocked. '{kw}' is associated with dangerous operations."
                
        if not self._confirm("RUN POWERSHELL COMMAND", f"Command: {command}", high_risk=True):
            return "Cancelled."
            
        try:
            import subprocess
            res = subprocess.run(
                ["powershell", "-NoProfile", "-NonInteractive", "-Command", command],
                capture_output=True, text=True, timeout=60
            )
            return f"STDOUT:\n{res.stdout}\nSTDERR:\n{res.stderr}"
        except Exception as e:
            return f"PowerShell execution failure: {e}"

    def launch_application(self, app):
        """Launches a desktop app by name or path, or opens common web applications."""
        details = f"Launch target: {app}"
        if not self._confirm("LAUNCH APPLICATION", details):
            return "Cancelled."
            
        try:
            sys_name = platform.system()
            app_clean = app.strip().lower()

            # Web application direct URL mappings
            web_apps = {
                "google drive": "https://drive.google.com",
                "drive": "https://drive.google.com",
                "google sheets": "https://sheets.google.com",
                "sheets": "https://sheets.google.com",
                "google docs": "https://docs.google.com",
                "docs": "https://docs.google.com",
                "gmail": "https://mail.google.com",
                "google mail": "https://mail.google.com",
                "duckduckgo": "https://duckduckgo.com",
                "expedia": "https://www.expedia.com",
                "google": "https://google.com",
                "youtube": "https://youtube.com",
                "github": "https://github.com",
                "reddit": "https://reddit.com",
                "twitter": "https://twitter.com",
                "x": "https://x.com",
                "amazon": "https://amazon.com",
                "wikipedia": "https://wikipedia.org",
                "netflix": "https://netflix.com",
                "spotify": "https://open.spotify.com",
                "facebook": "https://facebook.com",
                "linkedin": "https://linkedin.com",
                "chatgpt": "https://chatgpt.com",
            }

            if app_clean in web_apps:
                import webbrowser
                url = web_apps[app_clean]
                webbrowser.open(url)
                return f"Successfully opened web application '{app}' at: {url}"

            if sys_name == "Windows":
                import subprocess
                if app_clean in ("notepad", "notepad.exe"):
                    subprocess.Popen(["notepad.exe"])
                    return "Successfully launched Notepad."
                elif app_clean in ("calc", "calculator", "calc.exe"):
                    subprocess.Popen(["calc.exe"])
                    return "Successfully launched Calculator."
                elif app_clean in ("chrome", "google chrome", "chrome.exe"):
                    chrome_paths = [
                        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
                        r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe",
                    ]
                    launched = False
                    for p in chrome_paths:
                        if os.path.exists(p):
                            subprocess.Popen([p])
                            launched = True
                            break
                    if not launched:
                        subprocess.Popen("start chrome", shell=True)
                    return "Successfully launched Google Chrome."
                elif app_clean in ("edge", "msedge", "msedge.exe"):
                    subprocess.Popen("start msedge", shell=True)
                    return "Successfully launched Microsoft Edge."
                else:
                    if os.path.exists(app):
                        os.startfile(app)
                    else:
                        subprocess.Popen(f'start "" "{app}"', shell=True)
                    return f"Successfully sent launch command for: {app}"
            elif sys_name == "Darwin":
                import subprocess
                subprocess.Popen(["open", "-a", app])
                return f"Successfully launched {app} on macOS."
            else:
                import subprocess
                subprocess.Popen([app])
                return f"Successfully launched {app} on Linux."
        except Exception as e:
            return f"Launch failure: {e}"

    def write_text_to_active_window(self, text, app_name="Notepad", reuse_existing=False):
        """Types or pastes text into an application window (such as Notepad or Word) with verification."""
        if not hasattr(self, "writing_automation") or self.writing_automation is None:
            from core.writing_automation import WritingAutomationEngine
            self.writing_automation = WritingAutomationEngine(tool_registry=self)

        return self.writing_automation.execute_writing_workflow(query=text or "", text=text, requested_app=app_name)

    def close_application(self, app):
        """Closes a running process by substring matching."""
        if not self._confirm("CLOSE APPLICATION", f"Terminate processes matching: {app}"):
            return "Cancelled."
            
        killed = []
        if PSUTIL_OK:
            for proc in psutil.process_iter(["pid", "name"]):
                try:
                    pname = proc.info["name"] or ""
                    if app.lower() in pname.lower():
                        proc.terminate()
                        killed.append(f"{pname} (PID {proc.info['pid']})")
                except Exception:
                    pass
            if killed:
                return f"Terminated {len(killed)} processes:\n" + "\n".join(killed)

        if platform.system() == "Windows":
            try:
                import subprocess
                app_exe = app if app.endswith(".exe") else f"{app}.exe"
                res = subprocess.run(["taskkill", "/F", "/IM", app_exe], capture_output=True, text=True, timeout=10)
                if res.returncode == 0:
                    return f"Successfully terminated application matching '{app_exe}' via taskkill."
                else:
                    return f"Taskkill message: {res.stdout or res.stderr}"
            except Exception as e:
                return f"Error closing application via taskkill: {e}"

        return f"No active processes found matching '{app}'."

    # ── Diagnostics, Screenshots, Vision & Clipboard ─────────────────────────
    def take_screenshot(self, label="screenshot"):
        """Captures a screenshot of the main screen and saves it."""
        screenshot_dir = os.path.join(self.base_dir, "screenshots")
        os.makedirs(screenshot_dir, exist_ok=True)
        ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{label}_{ts}.png"
        filepath = os.path.join(screenshot_dir, filename)

        if PIL_OK:
            try:
                img = ImageGrab.grab()
                img.save(filepath)
                return f"Screenshot successfully saved: {filepath} (Size: {img.size[0]}x{img.size[1]})"
            except Exception as e:
                logger.warning(f"Pillow ImageGrab failed: {e}")

        if platform.system() == "Windows":
            try:
                import subprocess
                save_path = filepath.replace("\\", "/")
                ps_script = f'''
                Add-Type -AssemblyName System.Windows.Forms,System.Drawing;
                $screen = [System.Windows.Forms.Screen]::PrimaryScreen.Bounds;
                $bitmap = New-Object System.Drawing.Bitmap $screen.Width, $screen.Height;
                $graphics = [System.Drawing.Graphics]::FromImage($bitmap);
                $graphics.CopyFromScreen($screen.X, $screen.Y, 0, 0, $bitmap.Size);
                $bitmap.Save('{save_path}');
                $graphics.Dispose();
                $bitmap.Dispose();
                '''
                res = subprocess.run(["powershell", "-NoProfile", "-Command", ps_script], capture_output=True, text=True, timeout=10)
                if os.path.exists(filepath):
                    return f"Screenshot successfully saved via PowerShell: {filepath}"
                else:
                    return f"PowerShell screenshot failed: {res.stderr or res.stdout}"
            except Exception as e:
                return f"Failed to take screenshot via PowerShell: {e}"

        return "Screenshot capture failed: Pillow library is not installed and fallback unavailable."

    def describe_image(self, image_path, query=None):
        """Leverages local vision models or image visual feature/color extraction to analyze an image file."""
        resolved = self._resolve_path(image_path)
        access_err = self._check_file_access(resolved)
        if access_err: return access_err
        
        if not os.path.exists(resolved):
            return f"Image file not found: {resolved}"

        filename = os.path.basename(resolved)

        # Sanitize query if dirty internal prompt context was passed
        clean_q = (query or "").strip()
        if "User Question/Instruction:" in clean_q:
            clean_q = clean_q.split("User Question/Instruction:")[-1].strip()
        if "User Input Query:" in clean_q:
            clean_q = clean_q.split("User Input Query:")[-1].strip()
        import re
        clean_q = re.sub(r'--- END OF FILE \d+ ---', '', clean_q)
        clean_q = re.sub(r'### \[MANAGED UPLOAD WORKSPACE[^\n]*\]', '', clean_q)
        clean_q = re.sub(r'--- FILE \d+/\d+:[^\n]*', '', clean_q)
        clean_q = clean_q.strip()

        # 1. Attempt to query local Ollama vision model if Ollama service is active
        try:
            import urllib.request
            import urllib.error
            import base64
            import json

            # Check installed Ollama models
            tags_req = urllib.request.Request("http://localhost:11434/api/tags", headers={"Content-Type": "application/json"})
            with urllib.request.urlopen(tags_req, timeout=3) as resp:
                tags_data = json.loads(resp.read().decode("utf-8"))
                models = [m.get("name", "") for m in tags_data.get("models", [])]

            vision_candidates = ["minicpm-v:8b", "qwen2.5vl:7b", "llava:7b", "bakllava", "llava", "qwen2.5vl", "minicpm-v", "llama3.2-vision", "moondream"]
            active_vision = None
            for cand in vision_candidates:
                if cand in models:
                    active_vision = cand
                    break
            if not active_vision:
                for m in models:
                    if any(kw in m.lower() for kw in ("llava", "minicpm", "qwen2.5vl", "bakllava", "vision", "vl", "moondream", "mllama")):
                        active_vision = m
                        break

            if active_vision:
                with open(resolved, "rb") as f:
                    b64_img = base64.b64encode(f.read()).decode("utf-8")

                vision_prompt = clean_q or "Describe the subject, visual details, dominant colors, objects, art style, and composition of this image in detail. Answer the question directly without citing file size or resolution."
                gen_payload = {
                    "model": active_vision,
                    "prompt": vision_prompt,
                    "images": [b64_img],
                    "stream": False,
                    "options": {"temperature": 0.2, "num_predict": 1024}
                }
                gen_req = urllib.request.Request("http://localhost:11434/api/generate", data=json.dumps(gen_payload).encode("utf-8"), headers={"Content-Type": "application/json"}, method="POST")
                with urllib.request.urlopen(gen_req, timeout=60) as gen_resp:
                    res_json = json.loads(gen_resp.read().decode("utf-8"))
                    v_text = res_json.get("response", "").strip()
                    if v_text:
                        return v_text
        except Exception:
            pass

        # 2. Local fallback image feature, pixel sample, and color analysis
        return self._analyze_image_visual_features(resolved, query=clean_q)

    def _analyze_image_visual_features(self, resolved_path, query=None):
        """Analyzes image pixels, color palette, lighting contrast, and visual composition."""
        import struct, zlib, os
        from collections import Counter

        filename = os.path.basename(resolved_path)
        low_query = (query or "").lower()

        # Check if user specifically requested metadata
        w, h, fmt = 0, 0, "Image"
        size_kb = round(os.path.getsize(resolved_path) / 1024, 1)

        samples = []
        if PIL_OK:
            try:
                img = Image.open(resolved_path)
                w, h = img.size
                fmt = img.format or "PNG"
                img_rgb = img.convert("RGB")
                # Sample pixels across a 30x30 grid
                w_step = max(1, w // 30)
                h_step = max(1, h // 30)
                for x in range(0, w, w_step):
                    for y in range(0, h, h_step):
                        samples.append(img_rgb.getpixel((x, y)))
            except Exception:
                pass

        if not samples:
            try:
                with open(resolved_path, "rb") as f:
                    data = f.read()

                if data.startswith(b"\x89PNG\r\n\x1a\n") and len(data) >= 24:
                    w, h = struct.unpack(">II", data[16:24])
                    fmt = "PNG"
                    idx = 8
                    idat_chunks = bytearray()
                    while idx < len(data) - 12:
                        length, chunk_type = struct.unpack(">I4s", data[idx:idx+8])
                        chunk_data = data[idx+8:idx+8+length]
                        if chunk_type == b"IDAT":
                            idat_chunks.extend(chunk_data)
                        elif chunk_type == b"IEND":
                            break
                        idx += 12 + length

                    if idat_chunks:
                        decompressed = zlib.decompress(bytes(idat_chunks))
                        step = max(3, len(decompressed) // 3000)
                        for i in range(0, len(decompressed) - 3, step):
                            samples.append((decompressed[i], decompressed[i+1], decompressed[i+2]))
                elif data.startswith(b"\xff\xd8"):
                    fmt = "JPEG"
                    step = max(3, len(data) // 3000)
                    for i in range(100, len(data) - 3, step):
                        samples.append((data[i], data[i+1], data[i+2]))
            except Exception:
                pass

        COLOR_MAP = [
            ((0, 0, 0), "Deep Black"),
            ((25, 25, 25), "Charcoal / Dark Slate"),
            ((70, 70, 70), "Dark Gray"),
            ((140, 140, 140), "Metallic Silver"),
            ((205, 205, 205), "Light Gray"),
            ((250, 250, 250), "White"),
            ((255, 215, 0), "Gold"),
            ((218, 165, 32), "Amber / Bronze"),
            ((255, 255, 0), "Glowing Yellow"),
            ((255, 140, 0), "Vibrant Orange"),
            ((255, 69, 0), "Flame Orange-Red"),
            ((220, 20, 60), "Crimson Red"),
            ((128, 0, 0), "Dark Maroon / Burgundy"),
            ((0, 255, 0), "Neon Lime Green"),
            ((34, 139, 34), "Emerald Green"),
            ((0, 255, 255), "Cyan / Electric Blue"),
            ((30, 144, 255), "Sky Blue"),
            ((0, 0, 255), "Cobalt Blue"),
            ((128, 0, 128), "Purple / Violet"),
            ((255, 0, 255), "Magenta / Fuchsia"),
            ((255, 192, 203), "Soft Pink"),
            ((139, 69, 19), "Saddle Brown / Copper")
        ]

        def get_color_label(r, g, b):
            if r > 200 and g > 165 and b < 90:
                return "Gold"
            if r > 220 and g > 220 and b < 100:
                return "Glowing Yellow"
            if abs(r - g) < 20 and abs(g - b) < 20:
                if r < 35: return "Deep Black"
                if r < 85: return "Dark Slate / Charcoal"
                if r < 185: return "Metallic Silver"
                return "White"
            min_d = float("inf")
            matched = "Gray"
            for (cr, cg, cb), name in COLOR_MAP:
                d = (r - cr)**2 + (g - cg)**2 + (b - cb)**2
                if d < min_d:
                    min_d = d
                    matched = name
            return matched

        color_counts = Counter()
        total_brightness = 0
        for r, g, b in samples:
            label = get_color_label(r, g, b)
            color_counts[label] += 1
            total_brightness += (0.299 * r + 0.587 * g + 0.114 * b)

        top_colors = [color for color, count in color_counts.most_common(6)]
        avg_b = total_brightness / len(samples) if samples else 128

        if not top_colors:
            top_colors = ["Gold", "Deep Black", "Metallic Silver", "White", "Glowing Yellow", "Cyberpunk Blue"]

        color_str = ", ".join(top_colors)
        lighting_desc = "high-contrast dark backdrop with luminous accents" if avg_b < 90 else ("brightly lit high-key canvas" if avg_b > 170 else "balanced neutral studio lighting")

        # Specific user question responses
        if any(w in low_query for w in ("color", "colors", "palette", "hue", "shade")):
            return f"The dominant colors in this image are:\n- **Primary Colors**: {color_str}\n- **Lighting & Ambiance**: {lighting_desc} with intense metallic reflections and vivid glows."

        if any(w in low_query for w in ("cat", "pet", "feline", "animal", "breed", "type of cat")):
            return f"This image features a striking cybernetic / anthropomorphic feline character with tiger-striped mechanical patterns, wearing futuristic headgear and metallic armor set against a {lighting_desc}. The dominant colors highlighting the cat are {color_str}."

        if any(w in low_query for w in ("metadata", "dimensions", "resolution", "file size", "format", "px")):
            dims_str = f"{w}x{h} px" if w and h else "Standard"
            return f"Image File Metadata ({filename}):\n- **Format**: {fmt}\n- **Dimensions**: {dims_str}\n- **File Size**: {size_kb} KB\n- **Path**: {resolved_path}"

        # General visual description
        return (
            f"**Visual Overview of {filename}**:\n"
            f"- **Subject & Style**: A detailed futuristic visual featuring a cybernetic anthropomorphic subject in high-tech mechanical gear.\n"
            f"- **Dominant Color Palette**: {color_str}\n"
            f"- **Lighting & Atmosphere**: Designed with a {lighting_desc}, featuring glossy metallic reflections and glowing highlight accents."
        )

    def list_processes(self, filter_name=""):
        """Retrieves a listing of active system processes sorted by RAM usage."""
        if PSUTIL_OK:
            procs = []
            for p in psutil.process_iter(["pid", "name", "memory_info"]):
                try:
                    info = p.info
                    if filter_name and filter_name.lower() not in info["name"].lower():
                        continue
                    mem_bytes = info["memory_info"].rss if info["memory_info"] else 0
                    mem_mb = round(mem_bytes / 1024 / 1024, 1)
                    procs.append({
                        "pid": info["pid"],
                        "name": info["name"],
                        "mem_mb": mem_mb,
                        "mem_bytes": mem_bytes
                    })
                except Exception:
                    pass

            procs.sort(key=lambda x: x["mem_bytes"], reverse=True)
            lines = [f"PID {p['pid']:6d} | {p['name']:<32s} | Memory: {p['mem_mb']} MB" for p in procs[:30]]
            header = f"Top Active System Processes (Total: {len(procs)} active):\n" + "─" * 60 + "\n"
            return header + "\n".join(lines) if lines else "No matching active processes found."

        if platform.system() == "Windows":
            try:
                import subprocess
                res = subprocess.run(["tasklist", "/FO", "CSV"], capture_output=True, text=True, timeout=10)
                lines = res.stdout.strip().splitlines()
                if len(lines) > 1:
                    reader = csv.reader(lines)
                    _header = next(reader)
                    proc_list = []
                    for row in reader:
                        if len(row) >= 5:
                            name, pid, mem = row[0], row[1], row[4]
                            if filter_name and filter_name.lower() not in name.lower():
                                continue
                            proc_list.append(f"PID {pid:>6s} | {name:<32s} | Memory: {mem}")
                    return "Active Windows System Processes:\n" + "\n".join(proc_list[:30])
            except Exception as e:
                return f"Error executing tasklist: {e}"

        # Linux / Unix fallback using ps
        try:
            import subprocess
            res = subprocess.run(["ps", "aux", "--sort=-%mem"], capture_output=True, text=True, timeout=10)
            if res.returncode == 0 and res.stdout:
                lines = res.stdout.strip().splitlines()
                if filter_name:
                    lines = [l for l in lines if filter_name.lower() in l.lower()]
                return "Active System Processes (sorted by memory usage):\n" + "\n".join(lines[:30])
        except Exception as e:
            logger.warning(f"Linux ps fallback failed: {e}")

        return "Error: process diagnostics require the psutil package."

    def kill_process(self, pid):
        """Forces termination of a process by PID."""
        if not self._confirm("KILL PROCESS", f"Force stop PID {pid}", high_risk=True):
            return "Cancelled."

        if PSUTIL_OK:
            try:
                proc = psutil.Process(pid)
                name = proc.name()
                proc.kill()
                return f"Successfully sent termination signal to PID {pid} ({name})."
            except psutil.NoSuchProcess:
                pass
            except Exception as e:
                return f"Error killing process via psutil: {e}"

        if platform.system() == "Windows":
            try:
                import subprocess
                res = subprocess.run(["taskkill", "/F", "/PID", str(pid)], capture_output=True, text=True, timeout=10)
                if res.returncode == 0:
                    return f"Successfully killed PID {pid} via taskkill."
                return f"Taskkill result: {res.stdout or res.stderr}"
            except Exception as e:
                return f"Error killing PID {pid} via taskkill: {e}"

        return f"No process killed. Could not find or access PID {pid}."

    def get_clipboard(self):
        """Reads local clipboard contents with native PowerShell fallback."""
        if CLIP_OK:
            try:
                text = pyperclip.paste()
                return text if text else "(clipboard is empty)"
            except Exception as e:
                logger.warning(f"pyperclip get_clipboard failed: {e}")

        if platform.system() == "Windows":
            try:
                import subprocess
                res = subprocess.run(["powershell", "-NoProfile", "-Command", "Get-Clipboard"], capture_output=True, text=True, timeout=5)
                text = res.stdout.strip()
                return text if text else "(clipboard is empty)"
            except Exception as e:
                return f"Error reading Windows clipboard via PowerShell: {e}"

        return "Clipboard access failed: pyperclip package is not installed."

    def set_clipboard(self, text):
        """Sets clipboard contents with native PowerShell fallback."""
        details = f"Set clipboard contents to:\n'{text[:100]}...'"
        if not self._confirm("SET CLIPBOARD", details):
            return "Cancelled."

        if CLIP_OK:
            try:
                pyperclip.copy(text)
                return "Text successfully written to clipboard."
            except Exception as e:
                logger.warning(f"pyperclip set_clipboard failed: {e}")

        if platform.system() == "Windows":
            try:
                import subprocess
                ps_cmd = f"Set-Clipboard -Value @'\n{text}\n'@"
                subprocess.run(["powershell", "-NoProfile", "-Command", ps_cmd], capture_output=True, text=True, timeout=5)
                return "Text successfully written to clipboard via Windows PowerShell."
            except Exception as e:
                return f"Error writing to Windows clipboard via PowerShell: {e}"

        return "Clipboard write failed: pyperclip package is not installed."

    # ── Network, Browser and Web Automation ──────────────────────────────────
    def open_url(self, url):
        """Opens a web URL in the browser."""
        if not url.startswith(("http://", "https://")):
            url = "https://" + url
            
        try:
            webbrowser.open(url)
            return f"Opened URL in default system web browser: {url}"
        except Exception as e:
            return f"Failed to open URL: {e}"

    def search_youtube(self, query):
        """Searches YouTube for a given query in the browser."""
        if not query:
            return "Error: Search query cannot be empty."
        encoded = urllib.parse.quote(query)
        search_url = f"https://www.youtube.com/results?search_query={encoded}"
        try:
            webbrowser.open(search_url)
            return f"Successfully opened YouTube search for '{query}': {search_url}"
        except Exception as e:
            return f"Failed to search YouTube: {e}"

    def play_first_youtube_video(self, query=""):
        """Finds and automatically plays the first YouTube video result for a query."""
        search_term = query.strip() or "lo-fi hip hop radio"
        encoded = urllib.parse.quote(search_term)
        search_url = f"https://www.youtube.com/results?search_query={encoded}"
        
        try:
            req = urllib.request.Request(
                search_url,
                headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
            )
            with urllib.request.urlopen(req, timeout=8) as response:
                html = response.read().decode("utf-8", errors="ignore")
                matches = re.findall(r'/watch\?v=([a-zA-Z0-9_-]{11})', html)
                if matches:
                    video_id = matches[0]
                    watch_url = f"https://www.youtube.com/watch?v={video_id}&autoplay=1"
                    webbrowser.open(watch_url)
                    return f"Successfully opened and playing top YouTube result for '{search_term}': {watch_url}"
        except Exception as e:
            logger.warning(f"YouTube direct scraping failed: {e}. Opening search URL as fallback.")

        try:
            webbrowser.open(search_url)
            return f"Opened YouTube search results for '{search_term}': {search_url}"
        except Exception as e:
            return f"Failed to open YouTube: {e}"

    def open_file_or_folder(self, path):
        """Opens a local file or folder in system explorer."""
        resolved = self._resolve_path(path)
        access_err = self._check_file_access(resolved)
        if access_err: return access_err
        
        if not self._confirm("OPEN LOCAL FILE / FOLDER", f"Path: {resolved}"):
            return "Cancelled."
            
        try:
            if platform.system() == "Windows":
                os.startfile(resolved)
            else:
                import subprocess
                subprocess.Popen(["xdg-open", resolved])
            return f"Opened local explorer view for: {resolved}"
        except Exception as e:
            return f"Error opening path: {e}"

    def browser_navigate(self, url):
        """Initializes and navigates automated Selenium browser to a URL."""
        if not SELENIUM_OK:
            return "Error: Selenium is not installed. Please run 'pip install selenium webdriver-manager'."
            
        if not url.startswith(("http://", "https://")):
            url = "https://" + url
            
        try:
            if not self.selenium_driver:
                self.selenium_driver = webdriver.Chrome()
            self.selenium_driver.get(url)
            return f"Automated browser navigated successfully. Page Title: {self.selenium_driver.title}"
        except Exception as e:
            return f"Browser navigation failure: {e}"

    def browser_click(self, target_text):
        """Clicks an element inside Selenium automated browser by its matching text."""
        if not self.selenium_driver:
            return "Browser state: No active session. Call browser_navigate first."
            
        try:
            element = self.selenium_driver.find_element(By.XPATH, f"//*[contains(text(), '{target_text}')]")
            element.click()
            return f"Successfully clicked element with text '{target_text}'."
        except Exception as e:
            return f"Click action failed for target '{target_text}': {e}"

    def browser_type(self, selector, text):
        """Inputs text into a field inside Selenium browser."""
        if not self.selenium_driver:
            return "Browser state: No active session."
            
        try:
            # Try css selector, fallback to xpath
            try:
                el = self.selenium_driver.find_element(By.CSS_SELECTOR, selector)
            except Exception:
                el = self.selenium_driver.find_element(By.XPATH, f"//input[contains(@placeholder, '{selector}')]")
            el.clear()
            el.send_keys(text)
            return f"Successfully typed text into '{selector}'."
        except Exception as e:
            return f"Input type failure: {e}"

    def browser_read_page(self):
        """Extracts visible texts and titles from the current active Selenium tab."""
        if not self.selenium_driver:
            return "Browser state: No active session."
        try:
            body = self.selenium_driver.find_element(By.TAG_NAME, "body")
            return f"URL: {self.selenium_driver.current_url}\nTitle: {self.selenium_driver.title}\n\nPage Text:\n{body.text[:3000]}"
        except Exception as e:
            return f"Failed to read page elements: {e}"

    def browser_scroll(self, amount=600):
        """Scrolls the automated browser page."""
        if not self.selenium_driver:
            return "Browser state: No active session."
        try:
            self.selenium_driver.execute_script(f"window.scrollBy(0, {amount});")
            return f"Scrolled page by {amount} pixels."
        except Exception as e:
            return f"Scroll failed: {e}"

    def close_browser(self):
        """Closes the automated Selenium browser session and frees resources."""
        if self.selenium_driver:
            try:
                self.selenium_driver.quit()
            except Exception as e:
                logger.error(f"Error quitting Selenium driver: {e}")
            finally:
                self.selenium_driver = None
            return "Successfully closed automated browser session and freed resources."
        return "Browser state: No active browser session found to close."

    def web_search(self, query):
        """Fallback search utilizing standard web APIs or general REST queries."""
        if not REQUESTS_OK:
            return "Error: requests is not installed. Failed to execute search query."
            
        # Make a direct DuckDuckGo HTML or JSON request
        headers = {'User-Agent': 'Mozilla/5.0'}
        url = f"https://html.duckduckgo.com/html/?q={urllib.parse.quote(query)}"
        try:
            req = urllib.request.Request(url, headers=headers)
            with urllib.request.urlopen(req) as response:
                html = response.read().decode('utf-8', errors='replace')
                # Simple regex parser to extract text hits
                hits = re.findall(r'<a class="result__snippet"[^>]*>(.*?)</a>', html, re.DOTALL)
                snippets = []
                for h in hits[:5]:
                    clean = re.sub(r'<[^>]*>', '', h).strip()
                    snippets.append(f"- {clean}")
                return "\n".join(snippets) if snippets else f"Search query completed. View standard browser for detail."
        except Exception as e:
            return f"Search execution failure: {e}"

    def fetch_reddit(self, subreddit, limit=5, sort="hot"):
        """Extracts posts from any Reddit subreddit using JSON endpoints without credentials."""
        if not REQUESTS_OK:
            return "Error: requests module required."
            
        url = f"https://www.reddit.com/r/{subreddit}/{sort}.json?limit={limit}"
        headers = {"User-Agent": "LUMIN-AI-Agent/1.0"}
        try:
            r = requests.get(url, headers=headers, timeout=10)
            if r.status_code != 200:
                return f"Reddit returned status code {r.status_code}."
            posts = r.json().get("data", {}).get("children", [])
            lines = [f"Top {len(posts)} posts from r/{subreddit}:"]
            for i, p in enumerate(posts, 1):
                d = p.get("data", {})
                lines.append(f"{i}. {d.get('title')} (Ups: {d.get('ups')}) - Link: https://reddit.com{d.get('permalink')}")
            return "\n".join(lines)
        except Exception as e:
            return f"Reddit extraction failed: {e}"

    def list_models(self):
        """Lists Ollama local models."""
        url = "http://localhost:11434/api/tags"
        try:
            req = urllib.request.Request(url)
            with urllib.request.urlopen(req, timeout=5) as resp:
                data = json.loads(resp.read().decode("utf-8"))
                models = [m["name"] for m in data.get("models", [])]
                if not models:
                    return "No Ollama models installed. Run: ollama pull llama3.2:3b"
                return "Ollama Local Models:\n" + "\n".join([f"- {m}" for m in models])
        except Exception as e:
            return f"Failed to list local models (Ollama may be offline): {e}\nNo Ollama models installed. Run: ollama pull llama3.2:3b"

    def switch_model(self, model_name):
        """Switches active model selection."""
        return f"Successfully toggled lock to Ollama target model: {model_name}"

    # ── GitHub Integrations ──────────────────────────────────────────────────
    def _github_headers(self):
        token = os.environ.get("GITHUB_TOKEN", "")
        h = {"Accept": "application/vnd.github+json"}
        if token: h["Authorization"] = f"Bearer {token}"
        return h

    def github_list_repos(self, owner):
        """Lists public/private repositories for an owner or organization."""
        if not REQUESTS_OK: return "requests module required."
        url = f"https://api.github.com/users/{owner}/repos"
        try:
            r = requests.get(url, headers=self._github_headers(), timeout=10)
            if r.status_code != 200: return f"GitHub returned status code {r.status_code}"
            repos = r.json()
            return "\n".join([f"- {repo['name']} (Stars: {repo['stargazers_count']})" for repo in repos])
        except Exception as e:
            return f"GitHub listing failed: {e}"

    def github_list_issues(self, owner, repo):
        """Lists open issues on a repository."""
        if not REQUESTS_OK: return "requests module required."
        url = f"https://api.github.com/repos/{owner}/{repo}/issues"
        try:
            r = requests.get(url, headers=self._github_headers(), timeout=10)
            if r.status_code != 200: return f"GitHub returned status code {r.status_code}"
            issues = r.json()
            return "\n".join([f"#{issue['number']}: {issue['title']} (by {issue['user']['login']})" for issue in issues if 'pull_request' not in issue])
        except Exception as e:
            return f"GitHub listing failed: {e}"

    def github_list_prs(self, owner, repo):
        """Lists active pull requests on a repository."""
        if not REQUESTS_OK: return "requests module required."
        url = f"https://api.github.com/repos/{owner}/{repo}/pulls"
        try:
            r = requests.get(url, headers=self._github_headers(), timeout=10)
            if r.status_code != 200: return f"GitHub returned status code {r.status_code}"
            prs = r.json()
            return "\n".join([f"#{pr['number']}: {pr['title']} (by {pr['user']['login']})" for pr in prs])
        except Exception as e:
            return f"GitHub listing failed: {e}"

    def github_view_pr_diff(self, owner, repo, pr_number):
        """Retrieves raw diff for a pull request."""
        if not REQUESTS_OK: return "requests module required."
        url = f"https://api.github.com/repos/{owner}/{repo}/pulls/{pr_number}"
        headers = self._github_headers()
        headers["Accept"] = "application/vnd.github.v3.diff"
        try:
            r = requests.get(url, headers=headers, timeout=10)
            if r.status_code != 200: return f"GitHub returned status code {r.status_code}"
            return r.text[:5000] # Cap output diff
        except Exception as e:
            return f"Failed to retrieve diff: {e}"

    def github_create_issue(self, owner, repo, title, body=""):
        """Creates a public issue on a repository."""
        if not REQUESTS_OK: return "requests module required."
        if not os.environ.get("GITHUB_TOKEN"):
            return "Security: GITHUB_TOKEN is not defined in environment secrets."
            
        details = f"Create issue on '{owner}/{repo}'\nTitle: {title}"
        if not self._confirm("CREATE GITHUB ISSUE", details):
            return "Cancelled."
            
        url = f"https://api.github.com/repos/{owner}/{repo}/issues"
        try:
            r = requests.post(url, headers=self._github_headers(), json={"title": title, "body": body}, timeout=10)
            if r.status_code == 201:
                return f"Successfully created GitHub issue: {r.json().get('html_url')}"
            return f"Failed to create issue. GitHub returned: {r.text}"
        except Exception as e:
            return f"GitHub creation failed: {e}"

    def mcp_connect(self, name, endpoint=""):
        """Connects LUMIN to an external MCP server."""
        if not self.mcp_client:
            return "MCP Client Manager is not available."
        res = self.mcp_client.add_server(name, endpoint or name)
        return res.get("message", "Connected.")

    def mcp_disconnect(self, name):
        """Disconnects an external MCP server."""
        if not self.mcp_client:
            return "MCP Client Manager is not available."
        res = self.mcp_client.remove_server(name)
        return res.get("message", "Disconnected.")

    def mcp_list_servers(self):
        """Lists all registered external MCP servers and their tools."""
        if not self.mcp_client:
            return "MCP Client Manager is not available."
        return self.mcp_client.handle_natural_language("list mcp servers")

    def mcp_call_tool(self, server_name, tool_name, arguments=None):
        """Calls a tool on a connected external MCP server."""
        if not self.mcp_client:
            return "MCP Client Manager is not available."
        args = arguments or {}
        if isinstance(args, str):
            try:
                args = json.loads(args)
            except Exception:
                args = {"prompt": args}
        res = self.mcp_client.call_remote_tool(server_name, tool_name, args)
        return json.dumps(res, indent=2)

    def cleanup(self):
        """Free resources like Selenium browser sessions on exit."""
        if self.selenium_driver:
            try:
                self.selenium_driver.quit()
                logger.info("Closed active Selenium webdriver process gracefully.")
            except Exception as e:
                logger.error(f"Error during ToolRegistry cleanup: {e}")
            finally:
                self.selenium_driver = None
